"""
FRD Generator Service

Generation strategy:
1. Try Gemini first
2. If Gemini fails, try Hugging Face
3. If both fail, generate a strong deterministic fallback FRD

The service always returns a non-empty, professionally structured DOCX.
"""

from __future__ import annotations

import asyncio
import json
import logging
import re
from datetime import datetime
from html import unescape
from pathlib import Path
from typing import Any, Dict, List, Optional

import httpx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from huggingface_hub import InferenceClient

from app.config import settings

logger = logging.getLogger(__name__)


class FRDGeneratorService:
    SECTION_ORDER = [
        "overview",
        "document_history",
        "current_state",
        "proposed_solution",
        "roles",
        "application_types",
        "modules_and_applications",
        "process_flows",
        "functional_requirements",
        "non_functional_requirements",
        "integrations",
        "notifications",
        "reporting_visibility",
        "gap_analysis",
        "out_of_scope",
        "assumptions_constraints",
        "acceptance_signoff",
    ]

    def __init__(self) -> None:
        self.output_dir = Path(getattr(settings, "OUTPUT_DIR", "outputs"))
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.gemini_api_key = (getattr(settings, "GEMINI_API_KEY", "") or "").strip()
        self.gemini_model = (
            getattr(settings, "GEMINI_MODEL", "") or "gemini-2.0-flash"
        ).strip()

        self.hf_api_token = (getattr(settings, "HF_API_TOKEN", "") or "").strip()
        self.hf_model = (
            getattr(settings, "HF_MODEL", "") or "mistralai/Mistral-7B-Instruct-v0.3"
        ).strip()

        self.hf_client: Optional[InferenceClient] = None
        if self.hf_api_token:
            self.hf_client = InferenceClient(api_key=self.hf_api_token)

        # FIX: Track which model generated the FRD
        self._last_model_used: str = "Unknown"

    async def generate_frd(self, work_item_id: int, documents: List[Any]) -> Path:
        if not documents:
            raise ValueError("No source documents were provided for FRD generation.")

        normalized_docs = self._normalize_documents(documents)
        combined_source = self._combine_documents(normalized_docs)

        context = await self._extract_project_context(
            work_item_id=work_item_id,
            combined_source=combined_source,
            documents=normalized_docs,
        )

        sections: Dict[str, str] = {}
        for section_name in self.SECTION_ORDER:
            sections[section_name] = await self._generate_section(
                section_name=section_name,
                work_item_id=work_item_id,
                context=context,
                combined_source=combined_source,
            )

        output_path = self._build_docx(
            work_item_id=work_item_id,
            context=context,
            documents=normalized_docs,
            sections=sections,
        )

        logger.info(
            "FRD generated successfully | work_item_id=%s | model=%s | path=%s",
            work_item_id,
            self._last_model_used,
            output_path,
        )
        return output_path

    # -------------------------------------------------------------------------
    # Document preparation
    # -------------------------------------------------------------------------

    def _normalize_documents(self, documents: List[Any]) -> List[Dict[str, Any]]:
        normalized: List[Dict[str, Any]] = []

        for doc in documents:
            filename = getattr(doc, "filename", "unknown")
            content = getattr(doc, "content", "") or ""
            doc_type = getattr(doc, "doc_type", "other")
            url = getattr(doc, "url", None)

            clean_content = self._clean_text(content)
            clean_content = self._truncate(clean_content, 18000)

            normalized.append(
                {
                    "filename": filename,
                    "doc_type": doc_type,
                    "url": url,
                    "content": clean_content,
                }
            )

        normalized.sort(key=lambda d: self._doc_type_rank(d.get("doc_type", "other")))
        return normalized

    def _doc_type_rank(self, doc_type: str) -> int:
        order = {"sow": 0, "mom": 1, "transcript": 2, "other": 3}
        return order.get((doc_type or "other").lower(), 9)

    def _combine_documents(self, documents: List[Dict[str, Any]]) -> str:
        parts: List[str] = []

        for idx, doc in enumerate(documents, start=1):
            parts.append(
                f"""
===== SOURCE DOCUMENT {idx} =====
File Name: {doc["filename"]}
Document Type: {doc["doc_type"]}
Content:
{doc["content"]}
"""
            )

        return self._truncate("\n\n".join(parts).strip(), 50000)

    def _clean_text(self, text: str) -> str:
        text = unescape(text or "")
        text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
        text = re.sub(r"</p>", "\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<li[^>]*>", "- ", text, flags=re.IGNORECASE)
        text = re.sub(r"</li>", "\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _truncate(self, text: str, max_chars: int) -> str:
        text = text or ""
        if len(text) <= max_chars:
            return text
        return text[: max_chars - 200].rstrip() + "\n\n[TRUNCATED]"

    # -------------------------------------------------------------------------
    # Context extraction
    # -------------------------------------------------------------------------

    async def _extract_project_context(
        self,
        work_item_id: int,
        combined_source: str,
        documents: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        source_manifest = [
            {
                "filename": d["filename"],
                "doc_type": d["doc_type"],
                "chars": len(d["content"]),
            }
            for d in documents
        ]

        prompt = f"""
You are a senior business analyst creating an implementation-ready FRD context summary.

Return valid JSON only with this shape:
{{
  "project_name": "",
  "client_name": "",
  "business_context": "",
  "goals": ["", ""],
  "users": ["", ""],
  "roles": [
    {{
      "role": "",
      "responsibility": ""
    }}
  ],
  "application_types": ["", ""],
  "modules": [
    {{
      "module_name": "",
      "purpose": "",
      "users": ["", ""],
      "features": ["", ""],
      "key_fields": ["", ""],
      "validations": ["", ""]
    }}
  ],
  "integrations": [
    {{
      "system": "",
      "purpose": "",
      "data_exchanged": ""
    }}
  ],
  "notifications": ["", ""],
  "reporting_visibility": ["", ""],
  "out_of_scope": ["", ""],
  "assumptions_constraints": ["", ""],
  "gaps": [
    {{
      "gap": "",
      "proposed_solution": "",
      "reference_section": "",
      "phase": ""
    }}
  ]
}}

Rules:
1. Do not invent facts — only use what is in the source documents.
2. If a fact is unknown, use "To be confirmed".
3. Extract specific business details: client name, platform, process steps, user roles.
4. Keep the JSON concise but meaningful and specific to this project.

Work Item ID: {work_item_id}
Source Manifest:
{json.dumps(source_manifest, indent=2)}

Source Content:
{combined_source}
""".strip()

        raw = await self._call_model(prompt, max_output_tokens=2200, temperature=0.2)
        parsed = self._parse_json_response(raw)

        if parsed:
            return parsed

        logger.warning(
            "Structured context extraction failed; using deterministic fallback context."
        )
        return self._fallback_context(work_item_id, documents, combined_source)

    async def _generate_section(
        self,
        section_name: str,
        work_item_id: int,
        context: Dict[str, Any],
        combined_source: str,
    ) -> str:
        section_instructions = self._section_instructions()
        instruction = section_instructions[section_name]

        prompt = f"""
You are a senior functional consultant writing a polished, client-ready Functional Requirements Document.

Write only the requested section content.
Do not add a section title or heading.
Do not mention that you are an AI.
Do not invent facts — only use what is in the source documents.
Use "To be confirmed" only where information is genuinely unavailable.
Use specific names, systems, fields, and processes from the source documents.
Prefer implementation-ready language. Use markdown lists or markdown tables where appropriate.

Work Item ID: {work_item_id}

Structured Context (extracted from source documents):
{json.dumps(context, indent=2)}

Full Source Content:
{self._truncate(combined_source, 20000)}

Section Instructions:
{instruction}
""".strip()

        raw = await self._call_model(prompt, max_output_tokens=1600, temperature=0.25)
        cleaned = self._clean_llm_output(raw)

        if cleaned and len(cleaned) >= 40:
            return cleaned

        logger.warning(
            "LLM section generation failed; using fallback | section=%s", section_name
        )
        return self._fallback_section(section_name, context, combined_source)

    def _section_instructions(self) -> Dict[str, str]:
        return {
            "overview": """
Write a professional executive overview (3-4 paragraphs) covering:
- The specific business problem or opportunity this work item addresses
- The client name and their industry or business context
- What the proposed solution will do at a high level
- Measurable expected outcomes (e.g. reduce ticket resolution time, eliminate manual steps, improve visibility)
Be specific — use the actual client name, platform, and business context from the source documents.
No generic filler sentences. Every sentence must be grounded in the source material.
""",
            "document_history": """
Return a short markdown table:
| Date | Version | Description | Author |
Include today's date as version 1.0 drafted by FRD AI Agent.
""",
            "current_state": """
Based strictly on the source documents, describe:
1. How the process currently works (step by step where possible)
2. What systems or tools are currently used
3. Specific pain points, delays, manual steps, or errors mentioned
4. Who is affected and how
Be specific. Only state what is supported by the source material.
Do not use generic statements. Reference actual systems and teams from the documents.
""",
            "proposed_solution": """
Describe the future-state solution in concrete terms:
- What platform or technology will be used (be specific — e.g. Dynamics 365, Power Automate, Azure)
- What the core solution components are
- How it addresses each pain point from the current state
- What the user experience will look like
Use business language. Be specific to this project. Reference actual systems from the source.
""",
            "roles": """
Return a markdown table:
| Role | Responsibility |
List the specific roles involved in this solution with clear, implementation-ready responsibilities.
Use the actual role names and responsibilities from the source documents.
""",
            "application_types": """
Explain the specific solution components or application types required for this project.
Examples: Model-Driven App, Power Automate Flows, Customer Portal, API Layer, Admin Configuration UI, Power BI Reporting.
Only include what is grounded in the source documents.
""",
            "modules_and_applications": """
Break the solution into specific functional modules based on the source documents.
For each module use this format:

### [Module Name]
**Purpose:** [specific purpose]
**Users:** [specific user roles]
**Core Features:**
- [specific feature]
**Key Fields:**
- [field name]: [description and validation rules]
**Business Rules:**
- [specific rule]

Only include modules directly supported by the source documents.
Use real field names and business rules where mentioned.
""",
            "process_flows": """
Document the main end-to-end process flows as numbered steps.
For each major flow:
1. Name the flow clearly
2. List steps in format: [Actor] → [Action] → [System Response]
3. Include decision points and exception/error paths
Be specific to this project's actual described process, not generic steps.
""",
            "functional_requirements": """
Generate SPECIFIC, TESTABLE functional requirements directly derived from the source documents.
Each requirement must reference actual business functionality, not generic placeholders.

Format exactly as:
FR-001: The system shall [specific action] when [specific condition] so that [specific outcome].

Example of BAD requirement: "The solution shall capture business data."
Example of GOOD requirement: "The system shall automatically create a support ticket in Dynamics 365 when an inbound email is received at the configured support mailbox, capturing Subject, Sender Email, Body, and Received Timestamp as mandatory fields."

Generate at least 15 strong, specific requirements grounded in the source documents.
Each must be independently testable.
""",
            "non_functional_requirements": """
Generate non-functional requirements specific to this project's context.
Include actual numbers and specifics where possible:
- Performance: specific response time targets (e.g. "page load < 3 seconds for 95% of requests")
- Security: specific auth mechanism (e.g. "Azure AD SSO", "MFA enforced for all users")
- Compliance: specific standards mentioned in source (e.g. PCI DSS, GDPR, UAE data residency)
- Availability: specific SLA target (e.g. "99.5% uptime during business hours")
- Scalability: expected concurrent users or data volume
- Auditability: what must be logged and retained

Format: NFR-001: [Category] — [specific measurable requirement]
Generate at least 8 requirements.
""",
            "integrations": """
Describe all system integrations mentioned or clearly implied in the source documents.
Use this markdown table:
| System | Direction | Purpose | Data Exchanged | Trigger |

For each integration be specific about what data moves and when.
If details are unclear, state what needs to be confirmed and why it is needed.
""",
            "notifications": """
List specific notification events required by this solution:
- Trigger condition
- Recipient role
- Channel (email / SMS / in-app / Teams)
- Content summary
Base this on the actual workflow described in the source documents.
""",
            "reporting_visibility": """
List specific reports, dashboards, and visibility requirements:
- Report or dashboard name
- Purpose and intended audience
- Key metrics or fields displayed
- Refresh frequency if relevant
Be specific to this project's business needs.
""",
            "gap_analysis": """
Identify real gaps between what the source documents specify and what is needed for a complete implementation.
Use this markdown table:
| Gap | Impact | Proposed Resolution | Owner | Phase |

Only list genuine gaps found in the source material — missing field definitions, unclear business rules,
unconfirmed integrations, missing sign-off criteria, etc.
""",
            "out_of_scope": """
List items explicitly stated as out of scope in the source documents.
Also list items that are adjacent to the solution but not confirmed in scope.
Be specific — reference actual features, systems, or integrations.
""",
            "assumptions_constraints": """
**Assumptions:**
- List each assumption with its implication if the assumption proves incorrect

**Constraints:**
- Technical constraints (platform version, hosting region, compliance)
- Timeline or budget constraints mentioned
- Integration or third-party dependency constraints
""",
            "acceptance_signoff": """
Write a sign-off section including:
- Brief description of the review and approval process
- Sign-off table with roles specific to this project:
| Name | Role | Signature | Date |
""",
        }

    # -------------------------------------------------------------------------
    # Model calling chain
    # -------------------------------------------------------------------------

    async def _call_model(
        self,
        prompt: str,
        max_output_tokens: int = 1200,
        temperature: float = 0.2,
    ) -> str:
        # FIX: Try Gemini first with full error visibility
        if self.gemini_api_key:
            try:
                logger.info("Trying Gemini | model=%s", self.gemini_model)
                text = await self._call_gemini(
                    prompt=prompt,
                    max_output_tokens=max_output_tokens,
                    temperature=temperature,
                )
                if text and text.strip():
                    self._last_model_used = f"Gemini ({self.gemini_model})"
                    return text.strip()
            except Exception as exc:
                logger.warning("Gemini generation failed | error=%s", exc)

        # FIX: Try HuggingFace with async-safe call
        if self.hf_client and self.hf_api_token:
            try:
                logger.info("Trying HuggingFace | model=%s", self.hf_model)
                text = await self._call_huggingface(
                    prompt=prompt,
                    max_output_tokens=max_output_tokens,
                    temperature=temperature,
                )
                if text and text.strip():
                    self._last_model_used = f"HuggingFace ({self.hf_model})"
                    return text.strip()
            except Exception as exc:
                logger.warning("HuggingFace generation failed | error=%s", exc)

        # Both failed — deterministic fallback
        self._last_model_used = "Deterministic Fallback"
        logger.warning("All model providers failed — using deterministic fallback")
        return ""

    async def _call_gemini(
        self,
        prompt: str,
        max_output_tokens: int,
        temperature: float,
    ) -> str:
        # FIX: Use safe default model name
        model = self.gemini_model or "gemini-2.0-flash"
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"

        headers = {
            "Content-Type": "application/json",
            "x-goog-api-key": self.gemini_api_key,
        }

        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": max_output_tokens,
                "topP": 0.95,
            },
        }

        async with httpx.AsyncClient(timeout=120) as client:
            resp = await client.post(url, json=payload, headers=headers)

            # FIX: Log full error body so failures are visible in Azure logs
            if resp.status_code != 200:
                logger.error(
                    "Gemini API error | status=%s | body=%s",
                    resp.status_code,
                    resp.text[:1000],
                )
                resp.raise_for_status()

            data = resp.json()

        # FIX: Check for prompt blocking
        if data.get("promptFeedback", {}).get("blockReason"):
            logger.warning(
                "Gemini blocked prompt | reason=%s",
                data["promptFeedback"]["blockReason"],
            )
            return ""

        candidates = data.get("candidates", []) or []
        for candidate in candidates:
            # FIX: Check finish reason before using response
            finish_reason = candidate.get("finishReason", "")
            if finish_reason not in ("STOP", "MAX_TOKENS", ""):
                logger.warning("Gemini unexpected finishReason=%s", finish_reason)
                continue
            content = candidate.get("content", {}) or {}
            parts = content.get("parts", []) or []
            text_parts = [p.get("text", "") for p in parts if p.get("text")]
            if text_parts:
                result = "\n".join(text_parts).strip()
                logger.info("Gemini success | chars=%s", len(result))
                return result

        logger.warning(
            "Gemini returned empty candidates | response=%s", str(data)[:500]
        )
        return ""

    async def _call_huggingface(
        self,
        prompt: str,
        max_output_tokens: int,
        temperature: float,
    ) -> str:
        assert self.hf_client is not None

        messages = [
            {
                "role": "system",
                "content": (
                    "You are a senior business analyst writing implementation-ready "
                    "functional requirements documentation."
                ),
            },
            {"role": "user", "content": prompt},
        ]

        # FIX: Run blocking HuggingFace call in thread pool to avoid blocking async event loop
        def _sync_call():
            return self.hf_client.chat_completion(
                model=self.hf_model,
                messages=messages,
                max_tokens=max_output_tokens,
                temperature=max(temperature, 0.01),  # HF requires temperature > 0
            )

        loop = asyncio.get_event_loop()
        completion = await loop.run_in_executor(None, _sync_call)

        if completion and getattr(completion, "choices", None):
            message = completion.choices[0].message
            content = getattr(message, "content", None)
            if isinstance(content, str) and content.strip():
                logger.info("HuggingFace success | chars=%s", len(content))
                return content.strip()

        return ""

    def _parse_json_response(self, raw: str) -> Optional[Dict[str, Any]]:
        if not raw:
            return None

        raw = raw.strip()

        fenced = re.search(r"```(?:json)?\s*(\{.*\})\s*```", raw, re.DOTALL)
        if fenced:
            raw = fenced.group(1)

        try:
            parsed = json.loads(raw)
            if isinstance(parsed, dict):
                return parsed
        except Exception:
            pass

        start = raw.find("{")
        end = raw.rfind("}")
        if start >= 0 and end > start:
            try:
                parsed = json.loads(raw[start : end + 1])
                if isinstance(parsed, dict):
                    return parsed
            except Exception:
                return None

        return None

    def _clean_llm_output(self, text: str) -> str:
        text = (text or "").strip()
        if not text:
            return ""

        text = re.sub(r"^```(?:markdown|md|text)?\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s*```$", "", text)
        return text.strip()

    # -------------------------------------------------------------------------
    # Deterministic fallback generation
    # -------------------------------------------------------------------------

    def _fallback_context(
        self,
        work_item_id: int,
        documents: List[Dict[str, Any]],
        combined_source: str,
    ) -> Dict[str, Any]:
        project_name = self._infer_project_name(documents, combined_source)
        client_name = self._infer_client_name(combined_source)
        business_context = self._extract_relevant_paragraphs(
            combined_source,
            keywords=[
                "business",
                "problem",
                "current process",
                "manual",
                "pain point",
                "objective",
            ],
            default_text=(
                "The project aims to formalize and implement a target-state solution based on the "
                "available presales and discovery documents. Detailed business context should be "
                "validated during functional workshops."
            ),
            max_paragraphs=3,
        )

        goals = self._unique_non_empty(
            self._extract_bullets_by_keywords(
                combined_source,
                ["objective", "goal", "outcome", "benefit", "scope"],
                limit=5,
            )
        ) or [
            "Improve process efficiency and operational consistency.",
            "Reduce manual effort, ambiguity, and dependency on informal communication.",
            "Provide a scalable and auditable solution aligned with business needs.",
        ]

        users = self._unique_non_empty(
            self._extract_entities_by_keywords(
                combined_source,
                ["user", "team", "manager", "sales", "operations", "admin", "support"],
                limit=6,
            )
        ) or [
            "Business users",
            "Operations team",
            "Administrators",
        ]

        roles = [
            {
                "role": "Business User",
                "responsibility": "Initiates and performs day-to-day process steps in the solution.",
            },
            {
                "role": "Manager / Approver",
                "responsibility": "Reviews exceptions, approvals, and operational status.",
            },
            {
                "role": "System Administrator",
                "responsibility": "Maintains configuration, security, and reference data.",
            },
        ]

        modules = [
            {
                "module_name": "Intake and Data Capture",
                "purpose": "Capture required business inputs in a structured and validated format.",
                "users": ["Business User"],
                "features": [
                    "Structured form-based entry",
                    "Mandatory field validation",
                    "Status tracking",
                ],
                "key_fields": ["Primary identifiers", "Status", "Owner", "Dates"],
                "validations": [
                    "Mandatory fields",
                    "Format validation",
                    "Duplicate checks where applicable",
                ],
            },
            {
                "module_name": "Workflow and Processing",
                "purpose": "Route work through the required operational or approval flow.",
                "users": ["Business User", "Manager / Approver"],
                "features": [
                    "Stage-based progression",
                    "Business rule enforcement",
                    "Exception handling",
                ],
                "key_fields": ["Stage", "Assignment", "Decision outcome"],
                "validations": ["Transition rules", "Approval conditions"],
            },
            {
                "module_name": "Reporting and Visibility",
                "purpose": "Provide monitoring, auditability, and decision support.",
                "users": ["Manager / Approver", "System Administrator"],
                "features": [
                    "Operational dashboards",
                    "Search and filtering",
                    "Audit trail visibility",
                ],
                "key_fields": ["Status", "Owner", "Timestamps"],
                "validations": ["Role-based visibility"],
            },
        ]

        return {
            "project_name": project_name,
            "client_name": client_name,
            "business_context": business_context,
            "goals": goals,
            "users": users,
            "roles": roles,
            "application_types": [
                "Business application",
                "Workflow/automation layer",
                "Reporting/visibility layer",
            ],
            "modules": modules,
            "integrations": [
                {
                    "system": "To be confirmed",
                    "purpose": "Integration requirements to be finalized during detailed design.",
                    "data_exchanged": "To be confirmed",
                }
            ],
            "notifications": [
                "Status change notifications where business action is required.",
                "Exception or approval notifications for pending actions.",
            ],
            "reporting_visibility": [
                "Operational status visibility by work item or transaction stage.",
                "Management reporting for workload, turnaround time, and exceptions.",
                "Audit visibility into key updates and approvals.",
            ],
            "out_of_scope": [
                "Any features not explicitly supported by the approved scope documents.",
                "Downstream enhancements that require separate discovery or design approval.",
            ],
            "assumptions_constraints": [
                "Detailed field mapping and business rules will be confirmed during workshops.",
                "Access, security roles, and integrations depend on client environment readiness.",
            ],
            "gaps": [
                {
                    "gap": "Some functional details are not fully specified in the source documents.",
                    "proposed_solution": "Capture open points during requirement validation workshops and finalize in signed-off FRD.",
                    "reference_section": "4.5 Functional Requirements",
                    "phase": "Analysis",
                }
            ],
        }

    def _fallback_section(
        self,
        section_name: str,
        context: Dict[str, Any],
        combined_source: str,
    ) -> str:
        project_name = context.get("project_name") or "To be confirmed"
        client_name = context.get("client_name") or "To be confirmed"
        business_context = context.get("business_context") or "To be confirmed"

        if section_name == "overview":
            goals = context.get("goals") or []
            outcomes = "\n".join([f"- {g}" for g in goals]) or "- To be confirmed"
            return (
                f"The proposed initiative for **{project_name}** is intended to address the current "
                f"business need identified for **{client_name}**. Based on the available discovery and "
                f"presales material, the solution is expected to improve process control, reduce manual effort, "
                f"and provide better operational visibility.\n\n"
                f"**Business Context**\n{business_context}\n\n"
                f"**Expected Business Outcomes**\n{outcomes}"
            )

        if section_name == "document_history":
            today = datetime.now().strftime("%Y-%m-%d")
            return (
                "| Date | Version | Description | Author |\n"
                "|---|---|---|---|\n"
                f"| {today} | 1.0 | Initial draft generated from source documents | FRD AI Agent |"
            )

        if section_name == "current_state":
            return self._extract_relevant_paragraphs(
                combined_source,
                ["current", "manual", "pain", "issue", "challenge", "problem", "today"],
                (
                    "The current-state process appears to rely on manual coordination, fragmented information, "
                    "and limited visibility into progress and exceptions. These conditions may create delays, "
                    "inconsistencies, and operational dependency on individuals."
                ),
                max_paragraphs=4,
            )

        if section_name == "proposed_solution":
            return (
                "The proposed solution should provide a structured and governed business process supported by a "
                "centralized application layer, clearly defined workflow rules, and improved operational tracking. "
                "The future-state design should reduce ambiguity, standardize data capture, support exception handling, "
                "and provide reporting for business stakeholders."
            )

        if section_name == "roles":
            roles = context.get("roles") or []
            lines = ["| Role | Responsibility |", "|---|---|"]
            for role in roles:
                lines.append(
                    f"| {role.get('role', 'To be confirmed')} | {role.get('responsibility', 'To be confirmed')} |"
                )
            return "\n".join(lines)

        if section_name == "application_types":
            apps = context.get("application_types") or ["To be confirmed"]
            return "\n".join([f"- {item}" for item in apps])

        if section_name == "modules_and_applications":
            modules = context.get("modules") or []
            parts: List[str] = []
            for module in modules:
                parts.append(f"### {module.get('module_name', 'To be confirmed')}")
                parts.append(f"**Purpose:** {module.get('purpose', 'To be confirmed')}")
                parts.append("**Users:**")
                for user in module.get("users", []) or ["To be confirmed"]:
                    parts.append(f"- {user}")
                parts.append("**Core Features:**")
                for feat in module.get("features", []) or ["To be confirmed"]:
                    parts.append(f"- {feat}")
                parts.append("**Key Fields / Data Points:**")
                for field in module.get("key_fields", []) or ["To be confirmed"]:
                    parts.append(f"- {field}")
                parts.append("**Validations / Exceptions:**")
                for val in module.get("validations", []) or ["To be confirmed"]:
                    parts.append(f"- {val}")
                parts.append("")
            return "\n".join(parts).strip()

        if section_name == "process_flows":
            return (
                "1. A business user initiates a request or transaction by entering the required information.\n"
                "2. The system validates mandatory fields and applicable business rules.\n"
                "3. The record moves through the defined workflow stages and is routed to the appropriate owner or approver.\n"
                "4. Exceptions or validation failures are surfaced to the user for correction or review.\n"
                "5. Upon completion, the solution records the outcome and makes status/reporting information available to stakeholders."
            )

        if section_name == "functional_requirements":
            reqs = [
                "FR-001: The solution shall provide structured capture of business data required to initiate the process.",
                "FR-002: The solution shall validate mandatory fields before allowing progression to the next stage.",
                "FR-003: The solution shall maintain a status for each record throughout its lifecycle.",
                "FR-004: The solution shall support assignment or routing of records to appropriate users or teams.",
                "FR-005: The solution shall enforce business rules for stage transitions and exception handling.",
                "FR-006: The solution shall maintain an auditable history of key updates and actions.",
                "FR-007: The solution shall support search and retrieval of records using relevant business identifiers.",
                "FR-008: The solution shall provide role-based access to data and actions.",
                "FR-009: The solution shall support approval or review actions where required by the process.",
                "FR-010: The solution shall expose operational status and pending actions to relevant stakeholders.",
                "FR-011: The solution shall support notification triggers for significant business events.",
                "FR-012: The solution shall support reporting and visibility into workload, progress, and exceptions.",
            ]
            return "\n".join(reqs)

        if section_name == "non_functional_requirements":
            return "\n".join(
                [
                    "NFR-001: The solution shall enforce role-based access control for business data and actions.",
                    "NFR-002: The solution shall maintain an audit trail for create, update, approval, and status-change activities.",
                    "NFR-003: The solution shall provide acceptable response times for standard user actions under normal operating conditions.",
                    "NFR-004: The solution shall support maintainable configuration of reference data and business rules where feasible.",
                    "NFR-005: The solution shall ensure reliable processing and clear surfacing of validation or exception states.",
                    "NFR-006: The solution shall provide usable and consistent screens, forms, and messages for business users.",
                    "NFR-007: The solution shall support operational reporting and data visibility appropriate to user roles.",
                    "NFR-008: The solution shall align with client security, hosting, and compliance constraints to be confirmed during design.",
                ]
            )

        if section_name == "integrations":
            integrations = context.get("integrations") or []
            lines = ["| System | Purpose | Data Exchanged |", "|---|---|---|"]
            for item in integrations:
                lines.append(
                    f"| {item.get('system', 'To be confirmed')} | "
                    f"{item.get('purpose', 'To be confirmed')} | "
                    f"{item.get('data_exchanged', 'To be confirmed')} |"
                )
            return "\n".join(lines)

        if section_name == "notifications":
            notes = context.get("notifications") or ["To be confirmed"]
            return "\n".join([f"- {n}" for n in notes])

        if section_name == "reporting_visibility":
            items = context.get("reporting_visibility") or ["To be confirmed"]
            return "\n".join([f"- {n}" for n in items])

        if section_name == "gap_analysis":
            gaps = context.get("gaps") or []
            lines = [
                "| Gap | Proposed Solution | Reference Section | Phase |",
                "|---|---|---|---|",
            ]
            for item in gaps:
                lines.append(
                    f"| {item.get('gap', 'To be confirmed')} | "
                    f"{item.get('proposed_solution', 'To be confirmed')} | "
                    f"{item.get('reference_section', 'To be confirmed')} | "
                    f"{item.get('phase', 'To be confirmed')} |"
                )
            return "\n".join(lines)

        if section_name == "out_of_scope":
            items = context.get("out_of_scope") or ["To be confirmed"]
            return "\n".join([f"- {n}" for n in items])

        if section_name == "assumptions_constraints":
            items = context.get("assumptions_constraints") or ["To be confirmed"]
            return "\n".join([f"- {n}" for n in items])

        if section_name == "acceptance_signoff":
            return (
                "The final FRD shall be reviewed with business and project stakeholders. "
                "Any open points, assumptions, and unresolved decisions shall be captured before final sign-off.\n\n"
                "| Name | Role | Signature | Date |\n"
                "|---|---|---|---|\n"
                "| To be confirmed | Client Representative |  |  |\n"
                "| To be confirmed | Project Manager |  |  |\n"
                "| To be confirmed | Business Analyst |  |  |"
            )

        return "To be confirmed."

    def _infer_project_name(
        self, documents: List[Dict[str, Any]], combined_source: str
    ) -> str:
        for doc in documents:
            filename = doc.get("filename", "")
            stem = Path(filename).stem.strip()
            if stem and len(stem) > 3:
                return stem.replace("_", " ").replace("-", " ")

        patterns = [
            r"project\s*name\s*[:\-]\s*(.+)",
            r"engagement\s*name\s*[:\-]\s*(.+)",
            r"solution\s*name\s*[:\-]\s*(.+)",
        ]
        for pattern in patterns:
            match = re.search(pattern, combined_source, flags=re.IGNORECASE)
            if match:
                return match.group(1).strip()[:120]

        return "Work Item Project"

    def _infer_client_name(self, combined_source: str) -> str:
        patterns = [
            r"client\s*name\s*[:\-]\s*(.+)",
            r"customer\s*name\s*[:\-]\s*(.+)",
            r"company\s*name\s*[:\-]\s*(.+)",
        ]
        for pattern in patterns:
            match = re.search(pattern, combined_source, flags=re.IGNORECASE)
            if match:
                return match.group(1).strip()[:120]
        return "To be confirmed"

    def _extract_relevant_paragraphs(
        self,
        text: str,
        keywords: List[str],
        default_text: str,
        max_paragraphs: int = 3,
    ) -> str:
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        matches: List[str] = []

        for p in paragraphs:
            low = p.lower()
            if any(k.lower() in low for k in keywords):
                matches.append(p)
            if len(matches) >= max_paragraphs:
                break

        if matches:
            return "\n\n".join(matches)

        return default_text

    def _extract_bullets_by_keywords(
        self,
        text: str,
        keywords: List[str],
        limit: int = 5,
    ) -> List[str]:
        results: List[str] = []
        lines = [line.strip(" -\t") for line in text.splitlines() if line.strip()]
        for line in lines:
            low = line.lower()
            if any(k.lower() in low for k in keywords):
                results.append(line[:220])
            if len(results) >= limit:
                break
        return results

    def _extract_entities_by_keywords(
        self,
        text: str,
        keywords: List[str],
        limit: int = 6,
    ) -> List[str]:
        results: List[str] = []
        for line in [l.strip() for l in text.splitlines() if l.strip()]:
            low = line.lower()
            if any(k in low for k in keywords):
                results.append(line[:120])
            if len(results) >= limit:
                break
        return results

    def _unique_non_empty(self, items: List[str]) -> List[str]:
        seen = set()
        result = []
        for item in items:
            clean = (item or "").strip()
            key = clean.lower()
            if clean and key not in seen:
                seen.add(key)
                result.append(clean)
        return result

    # -------------------------------------------------------------------------
    # DOCX builder
    # -------------------------------------------------------------------------

    def _build_docx(
        self,
        work_item_id: int,
        context: Dict[str, Any],
        documents: List[Dict[str, Any]],
        sections: Dict[str, str],
    ) -> Path:
        doc = Document()
        self._set_doc_styles(doc)

        project_name = context.get("project_name") or f"Work Item {work_item_id}"
        client_name = context.get("client_name") or "To be confirmed"
        now = datetime.now()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("FUNCTIONAL REQUIREMENTS DOCUMENT")
        r.bold = True
        r.font.size = Pt(18)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(project_name)
        r.bold = True
        r.font.size = Pt(16)

        meta = doc.add_table(rows=0, cols=2)
        meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta.style = "Table Grid"

        for key, value in [
            ("Document Type", "Functional Requirements Document (FRD)"),
            ("Project Name", project_name),
            ("Client", client_name),
            ("Work Item ID", f"#{work_item_id}"),
            ("Version", "1.0 — Draft"),
            ("Status", "Auto-Generated — Pending Review"),
            # FIX: Show which model generated this FRD
            ("Generated By", f"FRD AI Agent — {self._last_model_used}"),
            ("Date", now.strftime("%B %d, %Y")),
        ]:
            row = meta.add_row().cells
            row[0].text = key
            row[1].text = str(value)

        doc.add_paragraph("")
        doc.add_heading("Source Documents", level=2)
        src_table = doc.add_table(rows=1, cols=2)
        src_table.style = "Table Grid"
        src_table.rows[0].cells[0].text = "File Name"
        src_table.rows[0].cells[1].text = "Type"

        for d in documents:
            row = src_table.add_row().cells
            row[0].text = d["filename"]
            row[1].text = str(d["doc_type"]).upper()

        doc.add_page_break()

        ordered_sections = [
            ("1. Overview", sections["overview"]),
            ("2. Document History", sections["document_history"]),
            ("3. Current State / Status Quo", sections["current_state"]),
            ("4. Proposed / Requested Solution", sections["proposed_solution"]),
            ("4.1 Roles", sections["roles"]),
            ("4.2 Type of Applications", sections["application_types"]),
            ("4.3 Modules and Applications", sections["modules_and_applications"]),
            ("4.4 Process Flows", sections["process_flows"]),
            ("4.5 Functional Requirements", sections["functional_requirements"]),
            (
                "4.6 Non-Functional Requirements",
                sections["non_functional_requirements"],
            ),
            ("4.7 Integrations", sections["integrations"]),
            ("4.8 Notifications", sections["notifications"]),
            ("4.9 Reporting / Visibility", sections["reporting_visibility"]),
            ("4.10 Gap Analysis", sections["gap_analysis"]),
            ("4.11 Out of Scope", sections["out_of_scope"]),
            ("4.12 Assumptions and Constraints", sections["assumptions_constraints"]),
            ("4.13 Acceptance / Sign-off", sections["acceptance_signoff"]),
        ]

        for title, content in ordered_sections:
            doc.add_heading(title, level=1 if re.match(r"^\d+\.", title) else 2)
            self._add_markdownish_content(doc, content)
            doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"FRD | {project_name} | WI#{work_item_id} | {self._last_model_used} | CONFIDENTIAL | {now.strftime('%Y-%m-%d')}"
        )
        run.italic = True
        run.font.size = Pt(9)

        output_path = (
            self.output_dir
            / f"FRD_WI{work_item_id}_{now.strftime('%Y%m%d_%H%M%S')}.docx"
        )
        doc.save(output_path)
        return output_path

    def _set_doc_styles(self, doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        styles = doc.styles
        if "Normal" in styles:
            styles["Normal"].font.name = "Calibri"
            styles["Normal"].font.size = Pt(11)

        if "Heading 1" in styles:
            styles["Heading 1"].font.name = "Calibri"
            styles["Heading 1"].font.size = Pt(14)
            styles["Heading 1"].font.bold = True

        if "Heading 2" in styles:
            styles["Heading 2"].font.name = "Calibri"
            styles["Heading 2"].font.size = Pt(12)
            styles["Heading 2"].font.bold = True

    def _add_markdownish_content(self, doc: Document, text: str) -> None:
        lines = [line.rstrip() for line in (text or "").splitlines()]
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            if (
                "|" in line
                and i + 1 < len(lines)
                and re.match(r"^\s*\|?[-:\s|]+\|?\s*$", lines[i + 1])
            ):
                table_lines = [line]
                i += 2
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_markdown_table(doc, table_lines)
                continue

            if re.match(r"^[-*]\s+", line):
                doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
                i += 1
                continue

            if re.match(r"^\d+\.\s+", line):
                doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
                i += 1
                continue

            if line.startswith("### "):
                doc.add_heading(line.replace("### ", "", 1), level=3)
                i += 1
                continue

            if line.startswith("## "):
                doc.add_heading(line.replace("## ", "", 1), level=2)
                i += 1
                continue

            if line.startswith("# "):
                doc.add_heading(line.replace("# ", "", 1), level=1)
                i += 1
                continue

            doc.add_paragraph(line)
            i += 1

    def _add_markdown_table(self, doc: Document, table_lines: List[str]) -> None:
        if not table_lines:
            return

        rows = []
        for line in table_lines:
            parts = [part.strip() for part in line.strip("|").split("|")]
            rows.append(parts)

        if len(rows) < 1:
            return

        header = rows[0]
        body = rows[1:]

        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Table Grid"

        for idx, value in enumerate(header):
            table.rows[0].cells[idx].text = value

        for row_data in body:
            if all(re.match(r"^[-:]+$", cell.replace(" ", "")) for cell in row_data):
                continue
            row = table.add_row().cells
            for idx in range(min(len(row_data), len(header))):
                row[idx].text = row_data[idx]
