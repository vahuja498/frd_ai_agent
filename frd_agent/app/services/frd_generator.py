"""
FRD Generator Service
Uses HuggingFace Inference API (free tier) to generate a professional
Functional Requirements Document from SOW, MOM, and Transcript documents.
"""

import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import List

import httpx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.webhook_payload import WorkItemDocument
from app.config import settings

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

# HuggingFace free Inference API endpoint
HF_API_URL = "https://api-inference.huggingface.co/models/{model}"

# Recommended free models (text generation / instruction following)
DEFAULT_MODEL = "mistralai/Mistral-7B-Instruct-v0.3"
FALLBACK_MODEL = "HuggingFaceH4/zephyr-7b-beta"


class FRDGeneratorService:
    def __init__(self):
        self.hf_token = settings.HF_API_TOKEN
        self.model = getattr(settings, "HF_MODEL", DEFAULT_MODEL)
        self._headers = {"Authorization": f"Bearer {self.hf_token}"}

    # ──────────────────────────────────────────────
    # Public API
    # ──────────────────────────────────────────────

    async def generate_frd(
        self,
        work_item_id: int,
        documents: List[WorkItemDocument],
        project_name: str = "Untitled Project",
        client_name: str = "",
    ) -> Path:
        """
        Orchestrates FRD generation:
        1. Build context from source documents
        2. Call HuggingFace LLM for each FRD section
        3. Render a professional .docx
        """
        context = self._build_document_context(documents)
        logger.info(f"Context length: {len(context)} chars")

        sections = await self._generate_frd_sections(context, work_item_id)
        docx_path = self._render_docx(
            work_item_id, sections, documents, project_name, client_name
        )
        return docx_path

    # ──────────────────────────────────────────────
    # Context Builder
    # ──────────────────────────────────────────────

    def _build_document_context(self, documents: List[WorkItemDocument]) -> str:
        """Concatenate document texts with clear section labels."""
        parts = []
        type_labels = {
            "sow": "📋 STATEMENT OF WORK (SOW)",
            "mom": "📝 MINUTES OF MEETING (MOM)",
            "transcript": "🎙️ MEETING TRANSCRIPT",
            "other": "📎 ADDITIONAL DOCUMENT",
        }
        for doc in documents:
            label = type_labels.get(doc.doc_type, "📎 DOCUMENT")
            # Truncate very large documents to stay within token limits
            content = doc.content[:8000] if len(doc.content) > 8000 else doc.content
            parts.append(f"=== {label}: {doc.filename} ===\n{content}\n")
        return "\n\n".join(parts)

    # ──────────────────────────────────────────────
    # LLM Section Generator
    # ──────────────────────────────────────────────

    async def _generate_frd_sections(self, context: str, work_item_id: int) -> dict:
        """Generate each FRD section via HuggingFace LLM."""
        section_prompts = {
            "project_overview": self._prompt_project_overview(context),
            "business_objectives": self._prompt_business_objectives(context),
            "scope": self._prompt_scope(context),
            "stakeholders": self._prompt_stakeholders(context),
            "functional_requirements": self._prompt_functional_requirements(context),
            "non_functional_requirements": self._prompt_nfr(context),
            "assumptions_constraints": self._prompt_assumptions(context),
            "risks": self._prompt_risks(context),
            "glossary": self._prompt_glossary(context),
        }

        results = {}
        for section_key, prompt in section_prompts.items():
            logger.info(f"Generating section: {section_key}")
            try:
                results[section_key] = await self._call_llm(prompt)
            except Exception as e:
                logger.warning(
                    f"LLM failed for section {section_key}: {e}. Using placeholder."
                )
                results[section_key] = (
                    f"[Content could not be generated. Please review source documents.]"
                )

        return results

    async def _call_llm(self, prompt: str) -> str:
        """Call HuggingFace Inference API."""
        url = HF_API_URL.format(model=self.model)
        payload = {
            "inputs": prompt,
            "parameters": {
                "max_new_tokens": 800,
                "temperature": 0.3,
                "do_sample": True,
                "return_full_text": False,
                "repetition_penalty": 1.1,
            },
        }

        async with httpx.AsyncClient(timeout=120) as client:
            resp = await client.post(url, headers=self._headers, json=payload)

            # Handle model loading (HF free tier may cold-start)
            if resp.status_code == 503:
                import asyncio

                logger.info("Model loading on HuggingFace, waiting 30s...")
                await asyncio.sleep(30)
                resp = await client.post(url, headers=self._headers, json=payload)

            resp.raise_for_status()
            data = resp.json()

        if isinstance(data, list) and data:
            text = data[0].get("generated_text", "")
        elif isinstance(data, dict):
            text = data.get("generated_text", "")
        else:
            text = str(data)

        return self._clean_llm_output(text)

    def _clean_llm_output(self, text: str) -> str:
        """Strip prompt echoes, leading/trailing whitespace, and markdown fences."""
        text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
        text = text.strip()
        # Remove any [INST] or <<SYS>> tokens if echoed
        text = re.sub(r"\[/?INST\]|<</?SYS>>", "", text)
        return text.strip()

    # ──────────────────────────────────────────────
    # Prompt Templates
    # ──────────────────────────────────────────────

    def _sys(self, role: str) -> str:
        return (
            f"<s>[INST] <<SYS>>\nYou are {role}. "
            "Write clearly, professionally, and concisely. "
            "Use bullet points for lists. Do not repeat the instructions.\n<</SYS>>\n\n"
        )

    def _prompt_project_overview(self, ctx: str) -> str:
        return (
            self._sys(
                "a senior business analyst writing a Functional Requirements Document"
            )
            + f"Based on the following source documents, write a professional PROJECT OVERVIEW section "
            f"for an FRD. Include: project name, background, purpose, and high-level description.\n\n"
            f"SOURCE DOCUMENTS:\n{ctx[:3000]}\n\nPROJECT OVERVIEW:[/INST]"
        )

    def _prompt_business_objectives(self, ctx: str) -> str:
        return (
            self._sys("a senior business analyst")
            + f"Extract and list the BUSINESS OBJECTIVES from these documents. "
            f"Each objective should be measurable and aligned with client goals.\n\n"
            f"SOURCE DOCUMENTS:\n{ctx[:3000]}\n\nBUSINESS OBJECTIVES:[/INST]"
        )

    def _prompt_scope(self, ctx: str) -> str:
        return (
            self._sys("a senior business analyst")
            + f"Write the SCOPE section for an FRD. Clearly define what is IN SCOPE and OUT OF SCOPE "
            f"based on the documents below.\n\n"
            f"SOURCE DOCUMENTS:\n{ctx[:3000]}\n\nSCOPE:[/INST]"
        )

    def _prompt_stakeholders(self, ctx: str) -> str:
        return (
            self._sys("a senior business analyst")
            + f"Identify all STAKEHOLDERS mentioned in the documents. "
            f"For each stakeholder provide: Name/Role, Organization, and Responsibility.\n\n"
            f"SOURCE DOCUMENTS:\n{ctx[:3000]}\n\nSTAKEHOLDERS:[/INST]"
        )

    def _prompt_functional_requirements(self, ctx: str) -> str:
        return (
            self._sys("a senior business analyst")
            + f"Extract and write detailed FUNCTIONAL REQUIREMENTS from the source documents. "
            f"Number each requirement (FR-001, FR-002, ...). Each requirement must have: "
            f"ID, Description, Priority (High/Medium/Low), and Source.\n\n"
            f"SOURCE DOCUMENTS:\n{ctx[:5000]}\n\nFUNCTIONAL REQUIREMENTS:[/INST]"
        )

    def _prompt_nfr(self, ctx: str) -> str:
        return (
            self._sys("a senior business analyst")
            + f"Write the NON-FUNCTIONAL REQUIREMENTS section. Cover: Performance, Security, "
            f"Scalability, Availability, Usability, and Compliance. Number as NFR-001, NFR-002, ...\n\n"
            f"SOURCE DOCUMENTS:\n{ctx[:3000]}\n\nNON-FUNCTIONAL REQUIREMENTS:[/INST]"
        )

    def _prompt_assumptions(self, ctx: str) -> str:
        return (
            self._sys("a senior business analyst")
            + f"List all ASSUMPTIONS AND CONSTRAINTS identified from the documents.\n\n"
            f"SOURCE DOCUMENTS:\n{ctx[:3000]}\n\nASSUMPTIONS AND CONSTRAINTS:[/INST]"
        )

    def _prompt_risks(self, ctx: str) -> str:
        return (
            self._sys("a senior business analyst")
            + f"Identify potential RISKS AND DEPENDENCIES from the source documents. "
            f"For each risk: Description, Likelihood, Impact, and Mitigation.\n\n"
            f"SOURCE DOCUMENTS:\n{ctx[:3000]}\n\nRISKS AND DEPENDENCIES:[/INST]"
        )

    def _prompt_glossary(self, ctx: str) -> str:
        return (
            self._sys("a senior business analyst")
            + f"Create a GLOSSARY of technical terms, abbreviations, and project-specific "
            f"terminology found in the documents.\n\n"
            f"SOURCE DOCUMENTS:\n{ctx[:2000]}\n\nGLOSSARY:[/INST]"
        )

    # ──────────────────────────────────────────────
    # DOCX Renderer
    # ──────────────────────────────────────────────

    def _render_docx(
        self,
        work_item_id: int,
        sections: dict,
        documents: List[WorkItemDocument],
        project_name: str = "Untitled Project",
        client_name: str = "",
    ) -> Path:
        """Renders a professional FRD Word document."""
        doc = Document()
        self._configure_document(doc)

        # Cover Page
        self._add_cover_page(doc, work_item_id, documents, project_name, client_name)

        # Table of Contents placeholder
        self._add_toc(doc)

        # FRD Sections
        section_map = [
            ("1. Project Overview", sections.get("project_overview", "")),
            ("2. Business Objectives", sections.get("business_objectives", "")),
            ("3. Scope", sections.get("scope", "")),
            ("4. Stakeholders", sections.get("stakeholders", "")),
            ("5. Functional Requirements", sections.get("functional_requirements", "")),
            (
                "6. Non-Functional Requirements",
                sections.get("non_functional_requirements", ""),
            ),
            (
                "7. Assumptions & Constraints",
                sections.get("assumptions_constraints", ""),
            ),
            ("8. Risks & Dependencies", sections.get("risks", "")),
            ("9. Glossary", sections.get("glossary", "")),
            ("10. Source Documents", self._format_source_list(documents)),
        ]

        for title, content in section_map:
            self._add_section(doc, title, content)

        # Footer
        self._add_footer(doc, work_item_id)

        output_path = (
            OUTPUT_DIR
            / f"FRD_WI{work_item_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        doc.save(output_path)
        logger.info(f"📄 FRD saved: {output_path}")
        return output_path

    def _configure_document(self, doc: Document):
        """Set document margins and default styles."""
        from docx.shared import Cm

        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    def _add_cover_page(
        self,
        doc: Document,
        work_item_id: int,
        documents: List[WorkItemDocument],
        project_name: str = "Untitled Project",
        client_name: str = "",
    ):
        """Create a professional cover page."""
        doc.add_paragraph()
        doc.add_paragraph()

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("FUNCTIONAL REQUIREMENTS DOCUMENT")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        doc.add_paragraph()

        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(project_name)
        sub_run.font.size = Pt(16)
        sub_run.font.bold = True
        sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

        doc.add_paragraph()

        meta = [
            ("Document Type", "Functional Requirements Document (FRD)"),
            ("Project Name", project_name),
            ("Client", client_name or "Not specified"),
            ("Work Item ID", f"#{work_item_id}" if work_item_id else "N/A"),
            ("Status", "DRAFT — Auto-Generated"),
            ("Generated By", "FRD AI Agent"),
            ("Date", datetime.now().strftime("%B %d, %Y")),
            (
                "Source Documents",
                ", ".join(
                    d.filename for d in documents if d.filename != "__metadata__.txt"
                )
                or "N/A",
            ),
        ]

        table = doc.add_table(rows=len(meta), cols=2)
        table.style = "Table Grid"
        for i, (key, value) in enumerate(meta):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = value
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_page_break()

    def _add_toc(self, doc: Document):
        """Add a Table of Contents placeholder."""
        doc.add_heading("Table of Contents", level=1)
        toc_para = doc.add_paragraph()
        toc_para.add_run(
            "Note: Update this Table of Contents in Microsoft Word by right-clicking → Update Field."
        ).italic = True
        doc.add_paragraph()

        toc_entries = [
            "1. Project Overview",
            "2. Business Objectives",
            "3. Scope",
            "4. Stakeholders",
            "5. Functional Requirements",
            "6. Non-Functional Requirements",
            "7. Assumptions & Constraints",
            "8. Risks & Dependencies",
            "9. Glossary",
            "10. Source Documents",
        ]
        for entry in toc_entries:
            p = doc.add_paragraph(entry, style="List Bullet")

        doc.add_page_break()

    def _add_section(self, doc: Document, title: str, content: str):
        """Add a numbered FRD section with heading and formatted content."""
        doc.add_heading(title, level=1)

        if not content or not content.strip():
            doc.add_paragraph("[No content generated for this section]").italic = True
        else:
            for line in content.strip().split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith(("- ", "• ", "* ")):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif re.match(r"^\d+\.", line):
                    doc.add_paragraph(line, style="List Number")
                elif line.isupper() and len(line) < 80:
                    sub = doc.add_paragraph()
                    sub.add_run(line).bold = True
                else:
                    doc.add_paragraph(line)

        doc.add_paragraph()

    def _add_footer(self, doc: Document, work_item_id: int):
        """Add a footer with document info."""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = (
            f"FRD — Work Item #{work_item_id} | "
            f"Auto-Generated by FRD AI Agent | "
            f"CONFIDENTIAL | {datetime.now().strftime('%Y-%m-%d')}"
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def _format_source_list(self, documents: List[WorkItemDocument]) -> str:
        if not documents:
            return "No source documents were attached."
        lines = []
        for doc in documents:
            lines.append(
                f"- {doc.filename} (Type: {doc.doc_type.upper()}, {len(doc.content)} chars extracted)"
            )
        return "\n".join(lines)
