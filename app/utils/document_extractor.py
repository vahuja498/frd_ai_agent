"""
Document Extractor Utility
Extracts plain text from .docx, .pdf, .txt, .md files.
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentExtractor:
    def extract_text(self, filename: str, content_bytes: bytes) -> str:
        ext = Path(filename).suffix.lower()

        try:
            if ext == ".docx":
                return self._extract_docx(content_bytes)
            elif ext == ".pdf":
                return self._extract_pdf(content_bytes)
            elif ext in (".txt", ".md", ".text"):
                return content_bytes.decode("utf-8", errors="replace")
            else:
                logger.warning(f"Unsupported file type: {ext}. Attempting raw decode.")
                return content_bytes.decode("utf-8", errors="replace")

        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}", exc_info=True)
            return f"[Failed to extract content from {filename}]"

    def _extract_docx(self, content_bytes: bytes) -> str:
        from docx import Document

        doc = Document(io.BytesIO(content_bytes))
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        result = "\n".join(paragraphs).strip()
        return result if result else "[DOCX contained no extractable text]"

    def _extract_pdf(self, content_bytes: bytes) -> str:
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
                pages = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        pages.append(text.strip())

                result = "\n\n".join(pages).strip()
                if result:
                    return result

        except ImportError:
            logger.warning("pdfplumber not installed. Trying PyPDF2...")
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}. Trying PyPDF2...")

        try:
            import PyPDF2

            reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
            texts = []

            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    texts.append(text.strip())

            result = "\n\n".join(texts).strip()
            return result if result else "[PDF contained no extractable text]"

        except ImportError:
            return "[PDF extraction failed: install pdfplumber or PyPDF2]"
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}", exc_info=True)
            return "[Failed to extract content from PDF]"
